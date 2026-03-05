from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Dict

from flask import Flask, jsonify, render_template, request, send_file
from werkzeug.utils import secure_filename

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import pytesseract
    from PIL import Image
except Exception:  # pragma: no cover
    pytesseract = None
    Image = None

from docx import Document

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024

ALLOWED_EXTENSIONS = {"pdf", "png", "jpg", "jpeg", "tiff", "bmp", "txt"}


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_invoice(file_storage) -> str:
    suffix = Path(file_storage.filename).suffix.lower()
    file_storage.stream.seek(0)

    if suffix == ".pdf":
        if PdfReader is None:
            raise RuntimeError("A PDF feldolgozáshoz telepíteni kell a pypdf csomagot.")
        reader = PdfReader(file_storage.stream)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    if suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        if pytesseract is None or Image is None:
            raise RuntimeError("A képfeldolgozáshoz telepíteni kell a pytesseract és pillow csomagokat.")
        image = Image.open(file_storage.stream)
        return pytesseract.image_to_string(image)

    if suffix == ".txt":
        return file_storage.stream.read().decode("utf-8", errors="ignore")

    raise RuntimeError("Nem támogatott fájltípus.")


def first_match(patterns: list[str], text: str) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            return match.group(1).strip()
    return "N/A"


def parse_invoice_fields(text: str) -> Dict[str, str]:
    normalized = re.sub(r"\r", "", text)
    fields = {
        "Számlaszám": first_match(
            [r"(?:Invoice\s*(?:No\.?|Number)|Bill\s*No\.?|No\.?)[\s:#-]*([A-Z0-9\-/]+)"],
            normalized,
        ),
        "Számla kelte": first_match(
            [r"(?:Invoice\s*Date|Date)\s*[:#-]?\s*([0-9]{1,2}[./-][0-9]{1,2}[./-][0-9]{2,4}|[A-Za-z]{3,9}\s+\d{1,2},?\s+\d{4})"],
            normalized,
        ),
        "Fizetési határidő": first_match(
            [r"(?:Due\s*Date|Payment\s*Due)\s*[:#-]?\s*([0-9]{1,2}[./-][0-9]{1,2}[./-][0-9]{2,4}|[A-Za-z]{3,9}\s+\d{1,2},?\s+\d{4})"],
            normalized,
        ),
        "Szállító neve": first_match(
            [r"(?:Seller|Supplier|Vendor|From)\s*[:#-]?\s*([^\n]{3,80})"],
            normalized,
        ),
        "Vevő neve": first_match(
            [r"(?:Buyer|Customer|Bill\s*To|To)\s*[:#-]?\s*([^\n]{3,80})"],
            normalized,
        ),
        "Nettó összeg": first_match(
            [r"(?:Subtotal|Net\s*Amount)\s*[:#-]?\s*([A-Z]{0,3}\s?[0-9., ]+)"],
            normalized,
        ),
        "ÁFA": first_match(
            [r"(?:VAT|Tax)\s*[:#-]?\s*([A-Z]{0,3}\s?[0-9., ]+|[0-9]{1,2}%?)"],
            normalized,
        ),
        "Bruttó végösszeg": first_match(
            [r"(?:Total\s*Due|Grand\s*Total|Total)\s*[:#-]?\s*([A-Z]{0,3}\s?[0-9., ]+)"],
            normalized,
        ),
        "Pénznem": first_match([r"\b(EUR|USD|GBP|CHF|PLN|CZK|RON|HUF)\b"], normalized),
    }
    return fields


def build_translation_doc(fields: Dict[str, str], source_text: str, original_filename: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Külföldi számla magyar fordítása", level=1)
    doc.add_paragraph(f"Forrásfájl: {original_filename}")
    doc.add_paragraph(f"Generálás dátuma: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Mező"
    hdr[1].text = "Kinyert adat"

    for label, value in fields.items():
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value or "N/A"

    doc.add_heading("Eredeti számlaszöveg (ellenőrzéshez)", level=2)
    snippet = source_text.strip() or "Nem sikerült szöveget kinyerni a feltöltött dokumentumból."
    doc.add_paragraph(snippet[:15000])

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


@app.get("/")
def index():
    return render_template("index.html")


@app.post("/api/translate-invoice")
def translate_invoice():
    if "invoice" not in request.files:
        return jsonify({"error": "Nem található feltöltött fájl (invoice)."}), 400

    file_storage = request.files["invoice"]
    if file_storage.filename == "":
        return jsonify({"error": "Nincs kiválasztott fájl."}), 400

    if not allowed_file(file_storage.filename):
        return jsonify({"error": "Csak PDF, kép vagy TXT fájl tölthető fel."}), 400

    try:
        extracted_text = extract_text_from_invoice(file_storage)
        fields = parse_invoice_fields(extracted_text)
        doc_stream = build_translation_doc(fields, extracted_text, secure_filename(file_storage.filename))

        output_filename = f"magyar_forditas_{Path(file_storage.filename).stem}.docx"
        return send_file(
            doc_stream,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
